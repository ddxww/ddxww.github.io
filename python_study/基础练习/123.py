from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()

def add_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(10.5)
content = """
Discussion (Each topic about 3 mins)

1. Pace of life in modern times

A: Hi, what do you think about the pace of life in modern society?
嗨，你觉得现代社会的生活节奏怎么样？

B: I think life today is much faster than before, especially for college students.
我觉得现在的生活比以前快多了，尤其是对大学生来说。

A: I agree. We have many things to deal with every day, such as classes, homework, exams, and club activities.
我同意。我们每天要处理很多事情，比如上课、作业、考试，还有社团活动。

B: Yes, sometimes I even feel that there is not enough time in a day.
是的，有时候我甚至觉得一天的时间根本不够用。

A: How does this fast pace affect you personally?
这种快节奏对你个人有什么影响？

B: It makes me feel stressed and tired, especially during exam weeks. I often stay up late to finish my work.
它让我感到压力和疲惫，尤其是在考试周。我经常熬夜完成任务。

A: Same here. Because of the fast pace, I sometimes don’t have time to exercise or relax.
我也是。因为节奏太快，我有时没时间锻炼或放松。

B: Do you prefer a fast pace or a slow pace of life?
你更喜欢快节奏还是慢节奏的生活？

A: I prefer a slower pace of life. I think life is not just about studying and working.
我更喜欢慢节奏的生活，我觉得生活不只是学习和工作。

B: Why do you think a slow pace is better?
你为什么觉得慢节奏更好？

A: A slower pace allows me to enjoy simple things, like having meals with friends or taking a walk on campus.
慢节奏能让我享受简单的小事，比如和朋友吃饭或在校园里散步。

B: That makes sense. I actually like a fast pace, because it gives me motivation to improve myself.
有道理。其实我更喜欢快节奏，因为它能激励我提升自己。

A: I think the most important thing is to find a balance between fast and slow.
我觉得最重要的是在快与慢之间找到平衡。

B: I totally agree. Balance can help us live a healthier life.
我完全同意，平衡能让我们生活得更健康。


2. A teacher who influenced you

A: Is there a teacher who has had a strong influence on you?
有没有一位对你影响很大的老师？

B: Yes, my high school English teacher influenced me a lot.
有的，我的高中英语老师对我影响很大。

A: What kind of person was she?
她是一个什么样的人？

B: She was very patient and always encouraged us to speak English in class.
她非常有耐心，总是鼓励我们在课堂上说英语。

A: Did she change your attitude toward learning?
她有没有改变你对学习的态度？

B: Yes. Before that, I was afraid of making mistakes, but she told us mistakes are part of learning.
有的。以前我害怕犯错，但她告诉我们犯错是学习的一部分。

A: That’s really meaningful.
这真的很有意义。

B: She also cared about our personal growth, not only our grades.
她不仅关心成绩，也关心我们的个人成长。

A: What did you learn from her?
你从她身上学到了什么？

B: I learned to be confident and to face challenges bravely.
我学会了自信，也学会勇敢面对挑战。

A: In your opinion, what are the most important qualities a teacher should have?
在你看来，老师最重要的品质是什么？

B: I think patience, responsibility, and encouragement are very important.
我觉得耐心、责任心和鼓励非常重要。

A: I agree. A good teacher should guide students and respect their differences.
我同意，好老师应该引导学生，并尊重学生的差异。

B: Teachers like that really leave a deep impression on students.
这样的老师真的会让学生终身难忘。


3. Environmental problems on campus

A: Have you noticed any environmental problems on our campus?
你有没有注意到校园里的环境问题？

B: Yes, I’ve noticed quite a few problems.
有的，我注意到不少问题。

A: What problems have you seen?
你看到了哪些问题？

B: For example, there is a lot of food waste in the cafeteria, and some students litter on campus.
比如食堂里浪费食物的现象很严重，还有学生乱扔垃圾。

A: I’ve seen that too, especially plastic bottles and snack bags.
我也见过，尤其是塑料瓶和零食袋。

B: These behaviors are harmful to the environment and also make the campus look messy.
这些行为既污染环境，也让校园看起来很乱。

A: What solutions do you suggest?
你有什么解决办法？

B: The cafeteria could offer smaller portions so students won’t waste food.
食堂可以提供小份菜，减少浪费。

A: That’s a good idea.
这是个好主意。

B: The school could also organize activities to raise students’ awareness of environmental protection.
学校也可以组织活动，提高学生的环保意识。

A: Students themselves should also take responsibility.
学生自己也应该承担责任。

B: Yes, protecting the environment needs everyone’s effort.
是的，环保需要大家共同努力。


4. AI: benefits and concerns

A: AI is developing very fast these days. What do you think about it?
现在人工智能发展得很快，你怎么看？

B: I think AI brings many benefits to our daily life.
我觉得人工智能给我们的生活带来了很多好处。

A: Can you give some examples?
你能举一些例子吗？

B: Sure. We can use AI to search for information, study more efficiently, and even plan our schedules.
当然。我们可以用 AI 查资料、更高效地学习，甚至安排日程。

A: That’s true. AI is really convenient.
确实如此，人工智能真的很方便。

B: However, there are also some concerns.
不过，也有一些担忧。

A: What kind of concerns?
什么样的担忧？

B: For example, privacy problems and job replacement.
比如隐私问题和工作被取代。

A: Do you think AI-powered robots will replace humans in the future?
你觉得 AI 机器人将来会取代人类吗？

B: I don’t think they will completely replace humans.
我觉得它们不会完全取代人类。

A: Why do you think so?
你为什么这么认为？

B: Because humans have creativity, emotions, and moral judgment, which AI doesn’t have.
因为人类有创造力、情感和道德判断，这是 AI 没有的。

A: I agree. AI should be used as a tool to help humans, not control them.
我同意，人工智能应该作为帮助人类的工具，而不是控制人类。

B: Learning how to use AI properly is very important for us.
学会正确使用人工智能对我们来说非常重要。


5. Becoming mature

A: Do you think physical age matters when it comes to maturity?
你觉得成熟和年龄有关吗？

B: I think age matters to some extent, but maturity is more about mindset and experience.
我觉得年龄有一定关系，但成熟更多取决于心态和经历。

A: Can you explain more?
你能具体说说吗？

B: Some young people are very responsible, while some older people may still be immature.
有些年轻人很有责任感，而有些年长的人却可能不够成熟。

A: That’s true.
确实如此。

B: Life experience helps people learn how to deal with problems calmly.
生活经历能帮助人们学会冷静地处理问题。

A: What advantages do young people have compared with older people?
你觉得年轻人相比年长的人有哪些优势？

B: Young people are energetic, open-minded, and willing to try new things.
年轻人更有活力、思想更开放，也更愿意尝试新事物。

A: I agree. We also learn new technologies faster.
我同意，而且我们学习新技术的速度更快。

B: What advice would you give to your future self?
你会给未来的自己什么建议？

A: I would tell myself to stay curious, be patient, and never stop learning.
我会告诉自己保持好奇、学会耐心、不断学习。

B: That’s good advice. I would also remind myself to care about my family and health.
这是很好的建议，我也会提醒自己多关心家人和健康。

A: Maturity is a lifelong process.
成熟是一个终身的过程。
"""
for line in content.split("\n"):
    add_para(line)

# ✅ 保存到当前项目目录
path = "Discussion_College_English_Full_3mins_each.docx"
doc.save(path)

print("Word 文件已成功生成：", path)
